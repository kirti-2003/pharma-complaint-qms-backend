from pathlib import Path

from docx import Document
from fastapi import HTTPException, status
from pypdf import PdfReader


def extract_text_from_file(
    storage_path: str,
    file_extension: str,
) -> str:
    """
    Extract readable text from TXT, DOCX, and PDF files.
    """

    file_path = Path(storage_path)

    if not file_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded file was not found.",
        )

    extension = file_extension.lower()

    try:
        if extension == ".txt":
            extracted_text = file_path.read_text(
                encoding="utf-8",
                errors="ignore",
            )

        elif extension == ".docx":
            document = Document(file_path)

            paragraphs = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]

            # Also extract text inside tables.
            table_rows = []

            for table in document.tables:
                for row in table.rows:
                    values = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]

                    if values:
                        table_rows.append(" | ".join(values))

            extracted_text = "\n".join(
                paragraphs + table_rows
            )

        elif extension == ".pdf":
            reader = PdfReader(str(file_path))

            extracted_text = "\n".join(
                page.extract_text() or ""
                for page in reader.pages
            )

        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=(
                    "Text extraction is currently supported only "
                    "for PDF, DOCX, and TXT files."
                ),
            )

    except HTTPException:
        raise

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Unable to read text from the uploaded document.",
        ) from exc

    extracted_text = extracted_text.strip()

    if not extracted_text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No readable text was found in the uploaded document.",
        )

    return extracted_text